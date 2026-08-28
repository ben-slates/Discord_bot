"""Private certificate generation for verified members."""

import asyncio
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import discord
from discord import app_commands
from discord.ext import commands

from database import SessionLocal, GuildConfig, VerificationRecord
from utils.db_executor import run_db


UUID_RE = re.compile(r"^0x[0-9A-Fa-f]{8}$")
TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "assets" / "ceritficate_template" / "input.docx"
CERTIFICATE_DIR = Path(__file__).resolve().parent.parent / "assets" / "batch1"


def _certificate_context(guild_id, user_id, verification_id):
    db = SessionLocal()
    try:
        config = db.query(GuildConfig).filter_by(guild_id=str(guild_id)).first()
        if not config or not config.certificate_enabled or not config.certificate_role or not config.certificate_channel:
            return None
        record = db.query(VerificationRecord).filter_by(
            verification_id=verification_id,
        ).first()
        if not record or int(record.discord_user_id) != int(user_id):
            return None
        return str(config.certificate_role), str(config.certificate_channel)
    finally:
        db.close()


def _certificate_config(guild_id):
    db = SessionLocal()
    try:
        config = db.query(GuildConfig).filter_by(guild_id=str(guild_id)).first()
        if not config or not config.certificate_enabled or not config.certificate_role or not config.certificate_channel:
            return None
        return str(config.certificate_role), str(config.certificate_channel)
    finally:
        db.close()


def _get_user_verification_id(user_id):
    db = SessionLocal()
    try:
        record = db.query(VerificationRecord).filter_by(discord_user_id=int(user_id)).first()
        return record.verification_id if record else None
    finally:
        db.close()


def _certificate_path(verification_id):
    return CERTIFICATE_DIR / f"certificate-{verification_id[2:].upper()}.pdf"


def _save_certificate_if_new(path, pdf_bytes):
    CERTIFICATE_DIR.mkdir(parents=True, exist_ok=True)
    try:
        with path.open("xb") as output:
            output.write(pdf_bytes)
        return True
    except FileExistsError:
        return False


def _replace_text(paragraph, replacements):
    for run in paragraph.runs:
        for placeholder, value in replacements.items():
            if placeholder in (run.text or ""):
                run.text = run.text.replace(placeholder, value)


def _render_certificate(name: str, team: str) -> bytes:
    """Render the supplied template to PDF; called outside the event loop."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Certificate generation requires the python-docx package.") from exc
    if not TEMPLATE_PATH.is_file():
        raise RuntimeError("Certificate template is missing from assets/ceritficate_template/input.docx.")
    soffice = shutil.which("soffice") or shutil.which("libreoffice") or shutil.which("lowriter")
    if not soffice:
        raise RuntimeError("Certificate PDF generation requires LibreOffice (soffice) on the server.")

    with tempfile.TemporaryDirectory(prefix="rynex-certificate-") as temp_dir:
        temp_path = Path(temp_dir)
        docx_path = temp_path / "certificate.docx"
        # The supplied template contains both the primary name and a smaller
        # repeated name line, plus the team placeholder.
        replacements = {"{name}": name, "{name_small}": name, "{team}": team}
        document = Document(str(TEMPLATE_PATH))
        for paragraph in document.paragraphs:
            _replace_text(paragraph, replacements)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_text(paragraph, replacements)
        document.save(str(docx_path))
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", str(docx_path), "--outdir", str(temp_path)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            raise RuntimeError("LibreOffice could not convert the certificate to PDF.") from exc
        pdf_path = temp_path / "certificate.pdf"
        if not pdf_path.is_file():
            raise RuntimeError("LibreOffice did not produce a certificate PDF.")
        return pdf_path.read_bytes()


class CertificateUUIDModal(discord.ui.Modal, title="Generate Certificate — UUID"):
    def __init__(self, cog):
        super().__init__()
        self.cog = cog
        self.uuid = discord.ui.TextInput(label="Verification ID", placeholder="0xXXXXXXXX", required=True, max_length=10)
        self.add_item(self.uuid)

    async def on_submit(self, interaction: discord.Interaction):
        raw_uuid = self.uuid.value.strip()
        verification_id = f"0x{raw_uuid[2:].upper()}" if raw_uuid[:2].lower() == "0x" else raw_uuid
        if not UUID_RE.fullmatch(verification_id):
            await interaction.response.send_message("Invalid verification ID format. Use your `0xXXXXXXXX` ID.", ephemeral=True)
            return
        # Discord only permits a modal as the response to an initial
        # interaction (not directly from another modal submission). Use a
        # private button to open the next modal step.
        await interaction.response.send_message(
            "UUID accepted. Click Continue to enter your name.",
            view=CertificateNameStepView(self.cog, verification_id),
            ephemeral=True,
        )


class CertificateNameStepButton(discord.ui.Button):
    def __init__(self, cog, verification_id):
        super().__init__(label="Continue", style=discord.ButtonStyle.primary)
        self.cog = cog
        self.verification_id = verification_id

    async def callback(self, interaction: discord.Interaction):
        await interaction.response.send_modal(CertificateNameModal(self.cog, self.verification_id))


class CertificateNameStepView(discord.ui.View):
    def __init__(self, cog, verification_id):
        super().__init__(timeout=300)
        self.add_item(CertificateNameStepButton(cog, verification_id))


class CertificateNameModal(discord.ui.Modal, title="Generate Certificate — Name"):
    def __init__(self, cog, verification_id):
        super().__init__()
        self.cog = cog
        self.verification_id = verification_id
        self.name = discord.ui.TextInput(label="Name", placeholder="Name printed on certificate", required=True, max_length=100)
        self.add_item(self.name)

    async def on_submit(self, interaction: discord.Interaction):
        await interaction.response.send_message(
            "Select your team:",
            view=CertificateTeamView(self.cog, self.verification_id, " ".join(self.name.value.split())),
            ephemeral=True,
        )


class CertificateTeamSelect(discord.ui.Select):
    def __init__(self, cog, verification_id, name):
        super().__init__(
            placeholder="Choose Red Team or Blue Team",
            min_values=1,
            max_values=1,
            options=[
                discord.SelectOption(label="Red Team", value="Red Team"),
                discord.SelectOption(label="Blue Team", value="Blue Team"),
            ],
        )
        self.cog = cog
        self.verification_id = verification_id
        self.certificate_name = name

    async def callback(self, interaction: discord.Interaction):
        await self.cog._complete_generation(
            interaction,
            self.verification_id,
            self.certificate_name,
            self.values[0],
        )


class CertificateTeamView(discord.ui.View):
    def __init__(self, cog, verification_id, name):
        super().__init__(timeout=300)
        self.add_item(CertificateTeamSelect(cog, verification_id, name))


class CertificateCog(commands.Cog):
    def __init__(self, bot):
        self.bot = bot

    @app_commands.command(name="generate-certificate", description="Generate your private certificate")
    async def generate_certificate(self, interaction: discord.Interaction):
        await interaction.response.send_modal(CertificateUUIDModal(self))

    async def _complete_generation(self, interaction, uuid, name, team):
        await interaction.response.defer(ephemeral=True)
        raw_uuid = uuid.strip()
        # Keep the required lowercase ``0x`` prefix while normalizing the
        # hexadecimal portion for case-insensitive user input.
        verification_id = f"0x{raw_uuid[2:].upper()}" if raw_uuid[:2].lower() == "0x" else raw_uuid
        certificate_config = await run_db(_certificate_config, interaction.guild_id)
        if not certificate_config:
            await interaction.followup.send("Certificate generation is not enabled for this server.", ephemeral=True)
            return
        if not UUID_RE.fullmatch(verification_id):
            await interaction.followup.send("Invalid verification ID format. Use your `0xXXXXXXXX` ID.", ephemeral=True)
            return
        if not name:
            await interaction.followup.send("Please provide a valid name.", ephemeral=True)
            return

        context = await run_db(
            _certificate_context,
            interaction.guild_id,
            interaction.user.id,
            verification_id,
        )
        if not context:
            await interaction.followup.send(
                "Your verification ID could not be confirmed, or it is not assigned to your account.",
                ephemeral=True,
            )
            return
        role_id, channel_id = certificate_config
        if str(interaction.channel_id) != channel_id:
            await interaction.followup.send("Use the configured certificate channel for this command.", ephemeral=True)
            return
        role = interaction.guild.get_role(int(role_id)) if interaction.guild else None
        if not role or role not in getattr(interaction.user, "roles", []):
            await interaction.followup.send("You do not have the required role to generate a certificate.", ephemeral=True)
            return

        certificate_path = _certificate_path(verification_id)
        if await asyncio.to_thread(certificate_path.is_file):
            await interaction.followup.send(
                "You already generated your certificate. Use `/get-critficate` to get your certificate.",
                ephemeral=True,
            )
            return

        try:
            pdf_bytes = await asyncio.to_thread(_render_certificate, name, team)
            saved = await asyncio.to_thread(_save_certificate_if_new, certificate_path, pdf_bytes)
        except Exception as exc:
            await interaction.followup.send(f"Certificate generation failed: {exc}", ephemeral=True)
            return
        if not saved:
            await interaction.followup.send(
                "You already generated your certificate. Use `/get-critficate` to get your certificate.",
                ephemeral=True,
            )
            return
        await interaction.followup.send(
            "Your certificate has been generated.",
            file=discord.File(str(certificate_path), filename=certificate_path.name),
            ephemeral=True,
        )

    @app_commands.command(name="get-critficate", description="Get your previously generated certificate")
    async def get_critficate(self, interaction: discord.Interaction):
        await interaction.response.defer(ephemeral=True)
        verification_id = await run_db(_get_user_verification_id, interaction.user.id)
        if not verification_id:
            await interaction.followup.send("You do not have a verification record yet.", ephemeral=True)
            return
        certificate_path = _certificate_path(verification_id)
        if not await asyncio.to_thread(certificate_path.is_file):
            await interaction.followup.send("You have not generated a certificate yet. Use `/generate-certificate`.", ephemeral=True)
            return
        await interaction.followup.send(
            "Here is your previously generated certificate.",
            file=discord.File(str(certificate_path), filename=certificate_path.name),
            ephemeral=True,
        )


async def setup(bot):
    await bot.add_cog(CertificateCog(bot))
